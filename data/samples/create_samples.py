import os
import sys

# Add backend to path for docx imports if needed
sys.path.append(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "backend"))

try:
    import docx
except ImportError:
    print("python-docx is not installed yet. Please run this script after backend packages are installed.")
    sys.exit(0)

SAMPLES_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(SAMPLES_DIR, exist_ok=True)

def create_john_doe_resume():
    """Generates John Doe's Data Science resume as a DOCX file."""
    doc = docx.Document()
    doc.add_heading("John Doe", 0)
    
    # Contact info
    p = doc.add_paragraph()
    p.add_run("john.doe@email.com | +1-555-0199 | San Francisco, CA\n").italic = True
    p.add_run("https://github.com/johndoe | https://linkedin.com/in/johndoe")
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Innovative Data Scientist with 3+ years of experience building machine learning models, scaling deep learning systems, and deploying analytical APIs. Passionate about solving business challenges using NLP, computer vision, and predictive modeling.")
    
    # Education
    doc.add_heading("Education", level=1)
    edu_p = doc.add_paragraph()
    edu_p.add_run("Master of Science in Computer Science (Concentration in AI)\n").bold = True
    edu_p.add_run("Stanford University, 2021 - 2023\n")
    edu_p.add_run("Bachelor of Science in Data Science\n").bold = True
    edu_p.add_run("University of California, Berkeley, 2017 - 2021")
    
    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Languages: Python, SQL, C++, R\n"
                    "Libraries: PyTorch, TensorFlow, Pandas, NumPy, Scikit-learn, Matplotlib, Seaborn\n"
                    "Tools & DevOps: Git, Docker, AWS (S3, EC2, Lambda), Linux, Bash, FastAPI\n"
                    "Databases: PostgreSQL, MongoDB, Redis")
    
    # Experience
    doc.add_heading("Experience", level=1)
    
    exp1 = doc.add_paragraph()
    exp1.add_run("Data Scientist - AI Innovations Corp (2023 - Present)\n").bold = True
    exp1.add_run("- Developed and deployed PyTorch deep learning models for NLP classification, improving accuracy by 15%.\n"
                 "- Containerized ML workflows using Docker and orchestrated training tasks on AWS EC2 nodes.\n"
                 "- Created high-performance FastAPI backends to serve real-time predictions with low latency.\n"
                 "- Wrote complex SQL queries to clean and analyze data from PostgreSQL databases.")
    
    exp2 = doc.add_paragraph()
    exp2.add_run("Machine Learning Intern - Analytics Tech (2022 - 2023)\n").bold = True
    exp2.add_run("- Built scikit-learn classification models to segment users based on platform activity.\n"
                 "- Designed interactive dashboards using Tableau and Matplotlib for executive reporting.\n"
                 "- Collaborated with software engineering team to integrate Python analytical pipelines.")
                 
    # Projects
    doc.add_heading("Projects", level=1)
    
    proj1 = doc.add_paragraph()
    proj1.add_run("Image Classification System (PyTorch & Docker)\n").bold = True
    proj1.add_run("- Trained CNN architectures on ImageNet dataset using PyTorch.\n"
                 "- Built a FastAPI web client and packaged the entire stack into a Docker container for deployment.")
                 
    proj2 = doc.add_paragraph()
    proj2.add_run("Interactive SQL Query Analyzer\n").bold = True
    proj2.add_run("- Created a Python web utility that parses and optimizes SQL queries using metadata statistics.")

    # Certifications
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("- AWS Certified Solutions Architect - Associate (2024)\n"
                    "- TensorFlow Developer Certificate (2022)")
                    
    doc.save(os.path.join(SAMPLES_DIR, "resume_john_doe.docx"))
    print("Generated resume_john_doe.docx")

def create_jane_smith_resume():
    """Generates Jane Smith's Frontend resume as a DOCX file."""
    doc = docx.Document()
    doc.add_heading("Jane Smith", 0)
    
    # Contact info
    p = doc.add_paragraph()
    p.add_run("jane.smith@email.com | +1-555-0144 | New York, NY\n").italic = True
    p.add_run("https://github.com/janesmith | https://janesmith.dev")
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Dynamic Frontend Developer with 4 years of experience building responsive, user-friendly web applications. Specializes in React, modern JavaScript, Tailwind CSS, and implementing premium UI animations.")
    
    # Education
    doc.add_heading("Education", level=1)
    edu_p = doc.add_paragraph()
    edu_p.add_run("Bachelor of Engineering in Software Engineering\n").bold = True
    edu_p.add_run("New York University, 2016 - 2020")
    
    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Frontend: JavaScript, TypeScript, React, Next.js, Redux, HTML5, CSS3, Tailwind CSS, Sass, Webpack, Vite\n"
                    "Backend: Node.js, Express, REST APIs\n"
                    "Tools & Testing: Git, GitHub Actions, Jest, Cypress, Figma, Agile/Scrum")
    
    # Experience
    doc.add_heading("Experience", level=1)
    
    exp1 = doc.add_paragraph()
    exp1.add_run("Senior Frontend Developer - PixelPerfect Labs (2022 - Present)\n").bold = True
    exp1.add_run("- Built and optimized customer dashboards using React, Redux Toolkit, and Vite, speeding up load times by 40%.\n"
                 "- Implemented a responsive design system using Tailwind CSS, ensuring pixel-perfect mobile-first compatibility.\n"
                 "- Configured CI/CD workflows using GitHub Actions to automate unit testing (Jest) and Vercel hosting deployment.\n"
                 "- Mentored junior frontend developers and conducted code reviews using Git version control.")
    
    exp2 = doc.add_paragraph()
    exp2.add_run("UI Developer - CodeCraft Solutions (2020 - 2022)\n").bold = True
    exp2.add_run("- Designed and implemented interactive features for multi-page web applications using React and Sass.\n"
                 "- Collaborated with design teams in Figma to translate mockups into semantic, clean HTML/CSS components.\n"
                 "- Built APIs with Node.js and Express to feed frontend states.")
                 
    # Projects
    doc.add_heading("Projects", level=1)
    
    proj1 = doc.add_paragraph()
    proj1.add_run("Modern Portfolio Generator (React & Tailwind)\n").bold = True
    proj1.add_run("- Developed an open-source react framework enabling developers to generate responsive portfolios with glassmorphism.\n"
                 "- Hosted on Vercel with automatic pull request preview deploys.")
                 
    proj2 = doc.add_paragraph()
    proj2.add_run("Interactive E-Commerce Checkout Flow\n").bold = True
    proj2.add_run("- Programmed a high-performance shopping cart checkout with complex client-state validation using Redux and Jest.")

    # Certifications
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("- Meta Front-End Developer Professional Certificate (2021)\n"
                    "- Scrum Alliance Certified ScrumMaster (CSM) (2023)")
                    
    doc.save(os.path.join(SAMPLES_DIR, "resume_jane_smith.docx"))
    print("Generated resume_jane_smith.docx")

def write_sample_jds():
    """Writes sample Job Descriptions as plain text files."""
    jd_ml = """Senior Machine Learning Engineer (Data Science & AI)
Role Overview:
We are looking for a Senior Machine Learning Engineer to join our AI team. You will build and scale deep learning systems, design NLP systems, and write RESTful prediction APIs.

Required Skills:
- Strong proficiency in Python, SQL, and Git
- Hands-on experience building neural networks with PyTorch
- Experience building and deploying models in production using Docker containers
- Experience designing backends with FastAPI
- Cloud deployment experience using AWS (EC2, S3, ECS)
- Strong understanding of data structures, algorithms, and system design
"""
    
    jd_web = """React & Frontend Engineer (Web Development)
Role Overview:
Join our design-focused team to build the next-generation dashboard interfaces. You will translate wireframes into beautiful, responsive web pages.

Required Skills:
- Professional experience with JavaScript, TypeScript, and React
- Advanced expertise in CSS3, HTML5, and Tailwind CSS
- Experience bundling code with Vite or Webpack
- Solid understanding of Git version control and GitHub workflows
- Experience building CI/CD pipelines with GitHub Actions
- Passion for animations, glassmorphism, and modern UI/UX design
"""

    with open(os.path.join(SAMPLES_DIR, "jd_senior_ml_engineer.txt"), "w", encoding="utf-8") as f:
        f.write(jd_ml.strip())
    print("Generated jd_senior_ml_engineer.txt")
    
    with open(os.path.join(SAMPLES_DIR, "jd_react_developer.txt"), "w", encoding="utf-8") as f:
        f.write(jd_web.strip())
    print("Generated jd_react_developer.txt")

if __name__ == "__main__":
    create_john_doe_resume()
    create_jane_smith_resume()
    write_sample_jds()
    print("All sample files generated under data/samples/ directory.")
